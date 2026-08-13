from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Planovany_strom_remesel_Remeslo_stredoveku.docx"
ASSET_DIR = ROOT / "assets" / "crafts"
TABLE_HELPERS = Path("/root/.codex/skills/builtins/documents/scripts")
sys.path.insert(0, str(TABLE_HELPERS))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x48)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x5E, 0x6B, 0x78)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
ACCENT_FILL = "EEF5EC"
CAUTION_FILL = "FFF8E8"
TABLE_WIDTH = 9360


def set_font(run, *, name="Calibri", size=None, color=TEXT, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_table_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_table(doc, headers, rows, widths, *, font_size=9.2, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    keep_table_row_together(hdr)
    for i, text in enumerate(headers):
        shade_cell(hdr.cells[i], header_fill)
        set_cell_text(hdr.cells[i], text, bold=True, color=DARK_BLUE, size=9.2)
    for row in rows:
        added_row = table.add_row()
        keep_table_row_together(added_row)
        cells = added_row.cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text), size=font_size)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Craft Bullet" not in styles:
        bullet = styles.add_style("Craft Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = styles["Craft Bullet"]
    bullet.base_style = normal
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(4)
    bullet.paragraph_format.line_spacing = 1.25
    bullet.paragraph_format.left_indent = Twips(540)
    bullet.paragraph_format.first_line_indent = Twips(-271)

    if "Craft Number" not in styles:
        number = styles.add_style("Craft Number", WD_STYLE_TYPE.PARAGRAPH)
    else:
        number = styles["Craft Number"]
    number.base_style = normal
    number.font.name = "Calibri"
    number.font.size = Pt(11)
    number.paragraph_format.space_before = Pt(0)
    number.paragraph_format.space_after = Pt(4)
    number.paragraph_format.line_spacing = 1.25
    number.paragraph_format.left_indent = Twips(540)
    number.paragraph_format.first_line_indent = Twips(-271)


def create_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs, default=0) + 1
    next_num = max(existing_num, default=0) + 1

    def add_definition(abstract_id, num_id, fmt, level_text, font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl_text_el = OxmlElement("w:lvlText")
        lvl_text_el.set(qn("w:val"), level_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, num_fmt, lvl_text_el, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            rpr.append(fonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", "Symbol")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    add_definition(next_abs + 2, next_num + 2, "decimal", "%1.")
    return next_num, (next_num + 1, next_num + 2)


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])


def add_bullets(doc, items, bullet_num):
    for item in items:
        p = doc.add_paragraph(style="Craft Bullet")
        apply_num(p, bullet_num)
        p.add_run(item)


def add_numbers(doc, items, decimal_num):
    for item in items:
        p = doc.add_paragraph(style="Craft Number")
        apply_num(p, decimal_num)
        p.add_run(item)


def add_callout(doc, title, body, *, fill=CALLOUT_FILL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Twips(120)
    p.paragraph_format.right_indent = Twips(120)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.keep_together = True
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), "C7D2E0")
        borders.append(border)
    ppr.append(borders)
    r = p.add_run(title)
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r.add_break()
    r2 = p.add_run(body)
    set_font(r2, size=10, color=TEXT)


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    r = p.add_run(label + ": ")
    set_font(r, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_font(r2)
    return p


def set_picture_alt_text(shape, *, title, description):
    """Add useful Word alt text to an inline picture."""
    doc_pr = shape._inline.docPr
    doc_pr.set("name", title)
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_illustration(doc, image_path, *, title, alt_text, width=1.15, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4 if caption is None else 1)
    p.paragraph_format.keep_together = True
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    set_picture_alt_text(shape, title=title, description=alt_text)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(5)
        cp.paragraph_format.keep_with_next = True
        run = cp.add_run(caption)
        set_font(run, size=8.5, color=MUTED, italic=True)
    return shape


def build_illustrated_side_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    keep_table_row_together(header)
    for cell, label in zip(header.cells, ("Pracovní dráha", "Činnosti, vazby a doporučený rozsah")):
        shade_cell(cell, HEADER_FILL)
        set_cell_text(cell, label, bold=True, color=DARK_BLUE, size=9.2)

    for row in rows:
        added = table.add_row()
        keep_table_row_together(added)
        left, right = added.cells
        left.text = ""
        picture_p = left.paragraphs[0]
        picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_p.paragraph_format.space_before = Pt(0)
        picture_p.paragraph_format.space_after = Pt(2)
        picture = picture_p.add_run().add_picture(str(ASSET_DIR / row["image"]), width=Inches(0.72))
        set_picture_alt_text(picture, title=f"Ilustrace dráhy {row['name']}", description=row["alt"])
        name_p = left.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_before = Pt(0)
        name_p.paragraph_format.space_after = Pt(0)
        set_font(name_p.add_run(row["name"]), size=9.0, color=DARK_BLUE, bold=True)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        right.text = ""
        info = right.paragraphs[0]
        info.paragraph_format.space_before = Pt(0)
        info.paragraph_format.space_after = Pt(0)
        info.paragraph_format.line_spacing = 1.08
        for index, (label, value) in enumerate(
            (("Činnosti", row["activities"]), ("Vazby", row["links"]), ("Rozsah", row["scope"]))
        ):
            if index:
                info.add_run().add_break()
            set_font(info.add_run(label + ": "), size=8.9, color=DARK_BLUE, bold=True)
            set_font(info.add_run(value), size=8.9, color=TEXT)
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    apply_table_geometry(
        table,
        [2050, 7310],
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=120,
        cell_margins_dxa={"top": 70, "bottom": 70, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


CRAFT_ICON_FILES = {
    "smithing": "main-00.png",
    "woodworking": "main-01.png",
    "masonry": "main-02.png",
    "wheelwrighting": "main-03.png",
    "pottery": "main-04.png",
    "leatherworking": "main-05.png",
    "textiles": "main-06.png",
    "bowyery": "main-07.png",
    "brewing": "main-08.png",
    "scribal": "main-09.png",
    "herbalism": "main-10.png",
    "healing": "main-11.png",
    "cooking": "main-12.png",
}


CRAFT_ICON_ALTS = {
    "smithing": "Kovadlina, kovářské kleště, kladivo a vykovaný hřebík.",
    "woodworking": "Sekera, dřevěná palice, dláto a čepový spoj.",
    "masonry": "Opracovaný pískovcový blok, kamenická palice a dláto.",
    "wheelwrighting": "Dřevěné loukoťové kolo, náboj a poříz koláře.",
    "pottery": "Hrnčířský kruh, hliněný džbán a připravená hlína.",
    "leatherworking": "Useň, šídlo, dratev a středověká obuv.",
    "textiles": "Tkalcovský stav, vřeteno, nůžky a složené lněné plátno.",
    "bowyery": "Rozpracovaný luk, šípové dříky, peří a pracovní nůž.",
    "brewing": "Dřevěná kádinka, sud, ječmen a chmel.",
    "scribal": "Brk, kalamář, pergamenové složky a svázaná kniha.",
    "herbalism": "Košík místních bylin a kamenný hmoždíř.",
    "healing": "Čisté obinadlo, dlahy, miska a nůžky pro ošetření rány.",
    "cooking": "Kotlík nad ohništěm, vařečka a zásoba sušených potravin.",
}


CRAFTS = [
    {
        "name": "Kovářství",
        "id": "smithing",
        "mentor": "Beneš, vesnický kovář; později hradní nebo městský mistr",
        "place": "Kovárna v Blatcích, hradní dílny, trh v Doksech",
        "mods": "síla, obratnost, soustředění, pozorování, odolnost; stav nářadí a výhně",
        "role": "Výroba a oprava nástrojů. Záměrně není postavené hlavně na zbraních; kovářství pohání zemědělství, dřevo, stavby i dopravu.",
        "spec": "Nástrojařství; podkovářství; stavební kování. Kolářské kování je mezioborový uzel s prací se dřevem.",
        "masterpiece": "Sada spolehlivých zemědělských nástrojů, kování mlýnského kola nebo kompletní kování vozu.",
        "rows": [
            ("Pozorovatel", "Bezpečnost výhně; druhy paliva; nástroje; barvy žáru; správné držení materiálu; práce s měchy.", "Udrží správný oheň, rozezná příliš studené a přepálené železo a bezpečně asistuje mistrovi."),
            ("Učeň", "Rovnání a broušení; hřebíky, skoby, háky; jednoduché opravy; základní kalení pod dohledem.", "Vyrobí stejnorodou sérii hřebíků a opraví srp bez praskliny nebo ztráty ostří."),
            ("Tovaryš", "Sekery, dláta, nože, motyky a podkovy; svařování ohněm; tepelné zpracování; odhad materiálu.", "Samostatně zhotoví nebo opraví celý pracovní nástroj a obhájí volbu postupu."),
            ("Specialista", "Volba specializace; složité spoje a kování; rozpoznání vad železa; přesnější řízení teploty.", "Zakázka se skrytou vadou: hráč musí problém objevit, napravit a neplýtvat materiálem."),
            ("Mistr", "Návrh postupu; hospodaření se železem a uhlím; diagnostika cizích výrobků; vedení společné práce.", "Mistrovský projekt přijme alespoň jeden další řemeslník, který na jeho spolehlivosti závisí."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje přípravky, spotřebu paliva nebo tvar nástroje. Mistr dílny učí, plánuje zásoby a ručí za jakost.", "Jedna cesta stačí k završení oboru; druhá zůstává dlouhodobým cílem."),
        ],
    },
    {
        "name": "Tesařství a truhlářství",
        "id": "woodworking",
        "mentor": "Matěj pro venkovské opravy; zkušený tesař v Doksech nebo na Housce",
        "place": "Statek, staveniště, hradní údržba, lesní pracovní místa",
        "mods": "obratnost, síla, pozorování, improvizace, zpracování dřeva, práce s lanem",
        "role": "Jedna společná větev do úrovně tovaryše; specialista se rozdělí na konstrukční tesařství, jemnější truhlářství nebo stavební opravy. Kolářství je kvůli geometrii kol a závislosti na kováři samostatný obor.",
        "spec": "Konstrukční tesařství; truhlářství a nábytek; stavební výdřeva a opravy. Bečvářství může být pozdější místní specializace, pokud bude mít region dostatečné odbytiště.",
        "masterpiece": "Nosná část krovu, přesná truhla s kováním nebo nové hradní dveře osazené do opraveného ostění.",
        "rows": [
            ("Pozorovatel", "Druhy dřeva; směr vláken; vlhkost; měření; bezpečné použití sekery, pily, pořízu a dláta.", "Vybere vhodné dřevo pro topůrko, kolík a nosný prvek a vysvětlí rozdíl."),
            ("Učeň", "Kolíky, topůrka, ploty, jednoduché spoje; ostření; oprava vrat a hospodářských předmětů.", "Opraví plot nebo vrata tak, aby spoj držel a dřevo zbytečně nepraskalo."),
            ("Tovaryš", "Čep a dlab; rámy; dveře; lavice, truhly a jednoduchý nábytek; čtení rozměrů z předlohy.", "Samostatně vytvoří předmět s několika přesnými spoji bez zakrytí chyb kovem nebo lanem."),
            ("Specialista", "Tesař řeší nosnost a stavbu; truhlář přesnost a povrch; stavební opravář diagnostiku starých konstrukcí.", "První zakázka, kde se chyba projeví až při sestavení; hráč musí najít příčinu."),
            ("Mistr", "Rozvržení materiálu; opravy starých konstrukcí; práce s křivým dřevem; vedení stavby nebo dílny.", "Mistrovský projekt obstojí při zatížení, počasí a kontrole objednavatele."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje spoj, pracovní přípravek nebo využití odpadu. Mistr dílny rozděluje práci, učí a ručí za stavbu.", "Završení oboru praktickým přínosem pro komunitu."),
        ],
    },
    {
        "name": "Hrnčířství",
        "id": "pottery",
        "mentor": "Hrnčíř v Dubé nebo Doksech; Zdislava pro léčebné nádoby",
        "place": "Hrnčířská dílna, naleziště hlíny, sušárna a pec",
        "mods": "obratnost, soustředění, pozorování, paměť; vlhkost hlíny, sušení a teplota pece",
        "role": "Zajišťuje nádoby pro vaření, skladování, léčiva i obchod. Výsledek vzniká v několika dnech a chyba se často projeví až při výpalu.",
        "spec": "Kuchyňské nádobí; zásobní a přepravní nádoby; kamnářské a stavební prvky.",
        "masterpiece": "Sada odolného kuchyňského nádobí, velká zásobnice nebo přesné kachle pro kamna.",
        "rows": [
            ("Pozorovatel", "Rozpoznání hlíny; čištění a hnětení; konzistence; části kruhu a pece; bezpečné sušení.", "Připraví hlínu bez kamínků a vzduchových kapes a vybere vhodné místo k sušení."),
            ("Učeň", "Ručně tvarované misky; jednoduché nádoby na kruhu; uchy; ořezání; základní výpal.", "Vyrobí několik nádob stejné velikosti, z nichž většina přežije sušení i výpal."),
            ("Tovaryš", "Tenkostěnné nádoby; poklice; džbány; engoba a jednoduchá výzdoba; skládání pece.", "Samostatná dávka pece s rozumnou zmetkovitostí a správným rozložením tepla."),
            ("Specialista", "Volba výrobkové větve; řízení atmosféry a teploty; náročné tvary; řešení prasklin a deformací.", "Rozpozná příčinu vady podle střepu a upraví příští dávku."),
            ("Mistr", "Návrh tvaru podle použití; plánování výroby; oprava pece; hospodaření s hlínou, palivem a časem.", "Mistrovská série obstojí při praktické zkoušce těsnosti, tepelného rázu nebo zatížení."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje směs, profil nádoby či vedení tepla. Mistr dílny řídí dávky, standardy a učedníky.", "Završení oboru opakovatelnou jakostí, ne jedním šťastným kusem."),
        ],
    },
    {
        "name": "Tkalcovství a krejčovství",
        "id": "textiles",
        "mentor": "Grete a ženy domácností; později tkalec, krejčí nebo barvíř v městském centru",
        "place": "Domácnost, pole lnu, máčírna, sušárna, stav a tržiště",
        "mods": "obratnost, soustředění, paměť, oprava oděvu, práce s lanem; čistota a kvalita vlákna",
        "role": "Propojuje pěstování lnu, zpracování vlákna, předení, tkaní, šití, barvení a opravy. První užitek přichází brzy, mistrovská látka až po dlouhém výrobním řetězci.",
        "spec": "Předení a tkaní; krejčovství a opravy; barvení a dokončování textilu. Provaznictví je kratší podpůrná dráha.",
        "masterpiece": "Souvislý kus rovnoměrného plátna, dobře padnoucí oděv nebo barevně a rozměrově stálá slavnostní textilie.",
        "rows": [
            ("Pozorovatel", "Celý řetězec lnu; třídění vláken; vřeteno, přeslice a stav; vady niti a látky.", "Správně seřadí kroky od sklizně po plátno a pozná nedostatečně zpracované vlákno."),
            ("Učeň", "Lámání a česání lnu; předení hrubé niti; jednoduché záplaty; příprava osnovy.", "Vytvoří použitelnou nit a opraví oděv bez dalšího trhání látky."),
            ("Tovaryš", "Rovnoměrná příze; osnova a útek; jednoduchá vazba; střih a sešití základního oděvu.", "Dokončí celý výrobek od vlákna po použitelný textilní předmět."),
            ("Specialista", "Volba textilní větve; jemnější vazby; přizpůsobení oděvu; moření, barvení a dokončení povrchu.", "Zakázka s omezeným materiálem prověří plánování střihu a schopnost opravit chybu."),
            ("Mistr", "Posouzení suroviny; návrh výrobku; barvení a dokončení; organizace domácí nebo dílenské výroby.", "Mistrovský výrobek musí být rovnoměrný, odolný a společensky přijatelný pro objednavatele."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje přípravek, vazbu nebo úsporu materiálu. Mistr dílny hlídá návaznost mnoha pracovníků a jakost.", "Završení oboru propojením techniky, organizace a vkusu."),
        ],
    },
    {
        "name": "Koželužství a ševcovství",
        "id": "leatherworking",
        "mentor": "Koželuh a švec v Dubé nebo Doksech; sedlář či řemenář podle specializace",
        "place": "Koželužna mimo hustou zástavbu, ševcovská dílna, stáje a trh",
        "mods": "obratnost, síla, odolnost, soustředění, hygiena; kvalita kůže, činidla a doba zpracování",
        "role": "Dlouhodobé zpracování surové kůže a následná výroba obuvi, opasků, brašen, pochev, řemení a postrojů.",
        "spec": "Ševcovství; řemenářství a brašnářství; sedlářství a postroje.",
        "masterpiece": "Pár dobře padnoucí obuvi, odolný postroj nebo kompletní sada řemení pro vůz či mlýn.",
        "rows": [
            ("Pozorovatel", "Druhy kůží; stažení a konzervace; čištění; činění; bezpečnost, zápach a hygiena pracoviště.", "Rozezná čerstvou, zkaženou a špatně činěnou kůži a připraví ji k dalšímu kroku."),
            ("Učeň", "Škrábání, napínání, mazání; řezání pásků; šití sedlářským stehem; jednoduché váčky a opasky.", "Zhotoví rovný šev a řemen, který se při zatížení neroztrhne."),
            ("Tovaryš", "Obuv, brašny, pochvy, závěsy a řemení; tvarování za mokra; opravy starých výrobků.", "Samostatný výrobek projde zkouškou pohybu, zatížení a vlhkosti."),
            ("Specialista", "Volba výrobkové větve; střih složitých dílů; hospodaření s plochou kůže; kombinace s textilem, dřevem a kovem.", "Hráč musí využít kůži s přirozenými vadami bez oslabení důležité části."),
            ("Mistr", "Posouzení celé usně; návrh podle těla zvířete či člověka; řízení činění; vedení zakázky.", "Mistrovský výrobek je pohodlný, opravitelný a dlouhodobě odolný."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje střih, spoj nebo impregnaci. Mistr dílny řídí zásoby, čistotu, zakázky a učedníky.", "Završení oboru spolehlivostí v každodenním používání."),
        ],
    },
    {
        "name": "Kamenictví a zednictví",
        "id": "masonry",
        "mentor": "Hradní kameník, zednická četa nebo mistr při opravách Housky",
        "place": "Lom, staveniště, vápenice, hradní zdi, kaple a hospodářské stavby",
        "mods": "síla, odolnost, obratnost, pozorování, soustředění; kvalita kamene, malty a založení",
        "role": "Od prostých oprav až po ostění, schody, klenby a pece. Hráč není osamělý stavitel hradu; vyšší práce jsou skupinové projekty.",
        "spec": "Zednictví; kamenické opracování; pece, krby a technické stavby.",
        "masterpiece": "Bezpečné ostění nebo schodiště, opravená část hradní zdi či dobře táhnoucí pec.",
        "rows": [
            ("Pozorovatel", "Druhy kamene; lomové vady; písek, vápno a malta; založení; přenos těžkých kusů a bezpečnost.", "Vybere vhodný kámen pro výplň, líc a opracovaný prvek a správně namíchá zkušební maltu."),
            ("Učeň", "Třídění kamene; jednoduché zdivo; spárování; dlažba; základní sekání a opravy.", "Postaví rovný úsek zdi, který drží vazbu a nevyboulí se."),
            ("Tovaryš", "Rohy, otvory, schody, sokly; přesnější opracování; lešení; čtení šablon.", "Dokončí stavební prvek v četě a správně naváže na cizí práci."),
            ("Specialista", "Volba větve; oblouky a ostění; komíny a pece; diagnostika trhlin, vlhkosti a sedání.", "Opraví závadu bez pouhého zakrytí její příčiny."),
            ("Mistr", "Rozvržení stavby; šablony; řízení čety; posouzení starého zdiva; plánování kamene, malty a času.", "Mistrovský projekt obstojí po zatížení, zatopení pece nebo nepřízni počasí."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje šablonu, tah pece nebo pracovní postup. Mistr dílny organizuje četu a ručí za bezpečnost.", "Završení oboru veřejně viditelnou a trvanlivou stavbou."),
        ],
    },
    {
        "name": "Bylinkářství",
        "id": "herbalism",
        "mentor": "Zdislava; dílčí znalosti mají Anna, děti, pasáci i další obyvatelé",
        "place": "Louky, cesty, potoky, lesní okraje, sušárna a bylinářčina domácnost",
        "mods": "sběr, pozorování, soustředění, paměť, vnímání, pohodlí v přírodě; roční období a stanoviště",
        "role": "Nejlépe připravená větev pro první plnou realizaci. Navazuje přímo na určování rostlin, volný popis znaků, sběr a zpracování.",
        "spec": "Poznávání a sběr; léčivé zpracování; lidové a obřadní použití. Poslední větev pracuje s vírou a nejistotou, ne s potvrzenými kouzly.",
        "masterpiece": "Spolehlivý herbář regionu a léčebná souprava pro konkrétní nemoc nebo zranění.",
        "rows": [
            ("Pozorovatel", "Tvar listu, květu a lodyhy; stanoviště; sezóna; bezpečné dotýkání; rozdíl mezi sběrem a určením.", "Správně popíše znaky běžné rostliny, i když ještě nezná její jméno."),
            ("Učeň", "Jitrocel, řebříček, kopřiva, šťovík a další základní druhy; správná část a čas sběru; čištění.", "Určí bezpečnou sadu běžných rostlin a nasbírá ji bez zničení stanoviště."),
            ("Tovaryš", "Jedovaté záměny; sušení; skladování; jednoduché čaje, obklady a masti ve spojení s ranhojičstvím.", "Připraví účinnou dávku a odhalí jednu nebezpečnou záměnu."),
            ("Specialista", "Volba větve; obtížné druhy; kombinace; sezónní plán; posuzování síly, plísně a znehodnocení.", "Vyřeší případ, kde správná rostlina nestačí bez vhodného sběru a zpracování."),
            ("Mistr", "Ucelený herbář; práce s nejistotou; kontraindikace; výuka znaků; zásoby pro domácnost nebo vesnici.", "Mistrovská zkouška zahrnuje určení, zpracování, dávku i vysvětlení rizik."),
            ("Inovátor / Mistr dílny", "Inovátor porovnává zkušenosti a zpřesňuje směsi. Mistr dílny vede sběr, sušení, záznamy a učedníky.", "Završení oboru znalostí, která je opakovatelná a bezpečně předatelná."),
        ],
    },
    {
        "name": "Ranhojičství",
        "id": "healing",
        "mentor": "Zdislava, zkušené ženy domácností a podle situace lazebník či městský ranhojič",
        "place": "Domácnosti, bylinářčina světnice, statek, cesta a místa pracovních úrazů",
        "mods": "pozorování, soustředění, paměť, empatie, odolnost vůči stresu; čistota, nástroje a stav pacienta",
        "role": "Praktická péče odpovídající možnostem kolem roku 1400. Neumožňuje zázračné léčení ani jistou diagnózu moderních nemocí.",
        "spec": "Rány a krvácení; kosti a pohybová zranění; horečky, otravy a dlouhodobá péče.",
        "masterpiece": "Zvládnutí vážného pracovního úrazu nebo epidemické situace včetně následné péče a zapojení komunity.",
        "rows": [
            ("Pozorovatel", "Posouzení vědomí, dýchání, krvácení a bolesti; hygiena; příprava vody, prádla, obvazů a prostoru.", "Správně určí pořadí pomoci a nezhorší stav zraněného."),
            ("Učeň", "Čištění ran; tlak a obvaz; drobné popáleniny; odpočinek, teplo a tekutiny; sledování zhoršení.", "Ošetří lehkou ránu a pozná, kdy musí přivolat zkušenějšího člověka."),
            ("Tovaryš", "Silnější krvácení; dlahy; horečka; běžné otravy; péče o nemocného; spolupráce s bylinkářstvím.", "Samostatně stabilizuje středně těžký stav a naplánuje další péči."),
            ("Specialista", "Volba větve; rozlišení příčin podle příznaků; práce pod časem; improvizace s omezenými prostředky.", "Případ s nejasnými příznaky prověří rozhodnutí, pozorování a ochotu přiznat nejistotu."),
            ("Mistr", "Třídění více zraněných; dlouhodobé sledování; příprava pomocníků; prevence nákazy a vedení záznamu.", "Mistrovská zkouška hodnotí přežití, komplikace i kvalitu následné péče, ne rychlost klikání."),
            ("Inovátor / Mistr dílny", "Inovátor zpřesňuje postupy pozorováním výsledků. Mistr dílny organizuje péči, zásoby a výuku.", "Završení oboru důvěryhodností a zodpovědností za druhé."),
        ],
    },
    {
        "name": "Vaření a uchování potravin",
        "id": "cooking",
        "mentor": "Anna a další hospodyně; mlynář pro mouku, rybáři a trh pro suroviny",
        "place": "Domácí ohniště, pec, mlýn, udírna, tržiště a komunitní slavnosti",
        "mods": "vaření, soustředění, paměť, improvizace, tolerance potravy; kvalita surovin, nádoba, oheň a hygiena",
        "role": "Základ přežití i přijetí do domácnosti. Od kaše a polévky vede k chlebu, kvašení, uzení a organizaci hostiny.",
        "spec": "Běžná kuchyně; pečení a práce s obilím; sušení, solení, uzení a skladování. Pivovarnictví je samostatný obor.",
        "masterpiece": "Bezpečné a dobře načasované jídlo pro celou domácnost nebo sváteční hostinu bez plýtvání zásobami.",
        "rows": [
            ("Pozorovatel", "Druhy ohniště; nádoby; voda a suroviny; známky zkažení; pořadí prací v domácí kuchyni.", "Připraví pracoviště, udrží oheň a pozná surovinu, která se nemá použít."),
            ("Učeň", "Kaše, jednoduché polévky, placky a bylinné nápoje; krájení, dávkování a řízení času.", "Uvaří jednoduché jídlo, které je bezpečné, syté a nepřipálené."),
            ("Tovaryš", "Chléb; složitější polévky; kombinace surovin; sušení, solení a základní kvašení; plán porcí.", "Připraví celé jídlo pro domácnost a správně načasuje více činností."),
            ("Specialista", "Volba větve; práce s těstem, masem, rybami nebo zásobami; náhrady chybějících surovin.", "Zakázka s omezenými zásobami prověří výživu, chuť, bezpečnost a hospodárnost."),
            ("Mistr", "Sezónní zásoby; velké dávky; slavnostní jídla; organizace pomocníků; diagnostika chyb pece a kvašení.", "Mistrovský projekt nasytí skupinu včas a bez následné nemoci či zbytečného odpadu."),
            ("Inovátor / Mistr dílny", "Inovátor vylepšuje postup z dostupných surovin. Mistr dílny řídí zásoby, oheň, pomocníky a výdej.", "Završení oboru schopností uživit komunitu, ne jen vytvořit drahý pokrm."),
        ],
    },
    {
        "name": "Kolářství",
        "id": "wheelwrighting",
        "mentor": "Zkušený kolář v Doksech nebo Dubé; Beneš dodává obruče a osové kování",
        "place": "Kolářská dílna, dvůr, sušárna dřeva, kovárna a cesty mezi Houskou, Doksy a Bezdězem",
        "mods": "obratnost, síla, pozorování, soustředění, zpracování dřeva; geometrie kola, vlhkost dřeva a přesnost kování",
        "role": "Samostatný obor na rozhraní truhlářství, dopravy a kovářství. Nejde jen o výrobu kola: hráč řeší náboje, loukotě, ráfky, nápravy, vůle a opravy vozu v terénu.",
        "spec": "Vozová kola; vozy a nápravy; velká vodní kola a převody ve spolupráci s mlynářem.",
        "masterpiece": "Pár shodných vozových kol s nasazenými obručemi nebo opravený vůz, který bezpečně obstojí na kamenité cestě.",
        "rows": [
            ("Pozorovatel", "Části kola a vozu; vhodné druhy dřeva; směr vláken; sušení; měření průměru, středu a vůle.", "Rozebere závadu kola na náboj, paprsek, loukoť, ráfek, obruč nebo nápravu a zvolí správný materiál."),
            ("Učeň", "Příprava šablon; dlabání náboje; opravy paprsků; mazání nápravy; jednoduché vozové díly.", "Opraví poškozený paprsek nebo vůli nápravy bez rozštípnutí okolního dřeva."),
            ("Tovaryš", "Sestavení kola; přesné rozteče; sesazení loukotí; spolupráce při nahřívání a nasazení železné obruče.", "Vyrobí kolo, které se otáčí bez výrazného házení a unese běžné zatížení."),
            ("Specialista", "Volba větve; těžké vozy; nápravy a převody; diagnostika nerovnoměrného opotřebení; nouzové opravy na cestě.", "Najde skrytou příčinu opakovaného praskání a neopraví pouze viditelný následek."),
            ("Mistr", "Návrh celého podvozku; párování kol; plánování dřeva a kování; vedení společné práce koláře a kováře.", "Mistrovský vůz nebo vodní kolo projde zatížením, delší jízdou či souvislým provozem."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje šablony, mazání, výměnnost dílů nebo rozložení zatížení. Mistr dílny řídí sušení, zásoby a návaznost na kovárnu.", "Završení oboru prokazatelně delší životností a opravitelností, nikoli nápadným vzhledem."),
        ],
    },
    {
        "name": "Výroba luků a šípů",
        "id": "bowyery",
        "mentor": "Lesník, zkušený lovec nebo hradní střelec; specializovaný výrobce na trhu či u posádky Bezdězu",
        "place": "Les pro výběr dřeva, suchá dílna, tržiště, hradní zázemí a střelnice pro ověření výrobku",
        "mods": "obratnost, pozorování, soustředění, paměť, zpracování dřeva; směr vláken, vlhkost, vyvážení a jakost šlachy",
        "role": "Výrobní větev oddělená od samotné lukostřelby. Odpovídá českému prostředí kolem roku 1400: lovecké a obranné luky, opravy, šípy a tětivy, nikoli anglická legenda o všudypřítomném dlouhém luku.",
        "spec": "Lukařství; šípařství; tětivy, pouzdra a polní opravy ve vazbě na textil a kůži.",
        "masterpiece": "Vyvážený lovecký luk s bezpečnou tětivou a sadou hmotnostně sladěných šípů, ověřený střelcem na různou vzdálenost.",
        "rows": [
            ("Pozorovatel", "Druhy dřeva; letokruhy a suky; části luku a šípu; bezpečné napínání; peří, hroty a tětivy.", "Vyřadí nebezpečný polotovar, určí směr vláken a rozezná poškozenou tětivu."),
            ("Učeň", "Rovnání dříků; zářezy; lepení a vázání per; jednoduché tětivy; údržba a drobné opravy.", "Vyrobí několik rovných šípových dříků se shodnou délkou a bezpečným končíkem."),
            ("Tovaryš", "Hrubování a postupné ohýbání luku; vyvažování ramen; osazení hrotů; hmotnostní třídění šípů.", "Dokončí běžný luk bez nebezpečného lomu a sadu šípů s předvídatelným letem."),
            ("Specialista", "Volba větve; kompenzace vlastností kusu dřeva; přesnější vyvažování; lovecké, cvičné a účelové šípy.", "Odhalí slabé místo při postupném napínání a výrobek opraví nebo bezpečně vyřadí."),
            ("Mistr", "Návrh podle postavy střelce a účelu; kontrola celé výrobní dávky; hospodaření s kvalitním dřevem a hroty.", "Mistrovská souprava je bezpečná, opravitelná a přijata zkušeným střelcem po praktické zkoušce."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje šablonu, měření průhybu či stálost šípů. Mistr dílny řídí sušení, párování dílů a kontrolu bezpečnosti.", "Završení oboru stabilitou celé soupravy, nikoli jedním mimořádným výstřelem."),
        ],
    },
    {
        "name": "Pivovarnictví",
        "id": "brewing",
        "mentor": "Zkušená hospodyně pro slabé domácí pivo; později sladovník nebo právem oprávněný pivovarník v Doksech či Dubé",
        "place": "Hospodářský dvůr, sladovna nebo městský pivovar; voda, hvozd, varná nádoba, kvasná káď a sklep",
        "mods": "soustředění, paměť, pozorování, vaření, hygiena a organizace práce; obilí, voda, teplota, kvasnice a čistota nádoby",
        "role": "Samostatný dlouhý proces od sladu přes vystírání a var po kvašení a skladování. Přístup závisí na místních právech a hospodářství; hráč si nepostaví volně pivovar v každé vsi.",
        "spec": "Sladování; domácí a běžné pivo; větší várky, skladování a provoz pivovaru podle místních práv.",
        "masterpiece": "Stálá várka pro svátek nebo pracovní četu, která je zdravotně bezpečná, chuťově vyrovnaná a připravená včas.",
        "rows": [
            ("Pozorovatel", "Obilí a slad; kvalita vody; nádoby; čistota; rozdíl mezi vystíráním, varem, chlazením a kvašením.", "Správně připraví a vyčistí zařízení, pozná zkaženou surovinu a udrží bezpečné pořadí práce."),
            ("Učeň", "Máčení a klíčení obilí; sušení sladu; drcení; ohřev vody; péče o malou kvasnou nádobu.", "Připraví použitelný slad a jednoduchou malou várku pod dohledem bez zkažení."),
            ("Tovaryš", "Vystírání; scezování; var; chmel či jiné dobové ochucení; chlazení, zakvašení a stáčení.", "Samostatně vede celou běžnou várku a včas rozpozná nesprávnou teplotu či začínající kontaminaci."),
            ("Specialista", "Volba větve; řízení síly a chuti; opakování stejné várky; využití mláta; diagnostika chyb sladu a kvašení.", "Zachrání ještě opravitelnou várku nebo ji včas vyřadí, aniž by ohrozil domácnost."),
            ("Mistr", "Plánování zásob, vody, paliva a nádob; větší várky; sklepní režim; organizace pomocníků a odbytu.", "Mistrovský nápoj obstojí při opakovaném uvaření a hodnocení více odběratelů."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje hospodaření s teplem, vodou nebo stálost várky. Mistr dílny řídí slad, var, kvašení, zásoby a čistotu.", "Završení oboru spolehlivým provozem v mezích dobových práv a dostupných surovin."),
        ],
    },
    {
        "name": "Písařství a knihařství",
        "id": "scribal",
        "mentor": "Kněz, písař hradní správy nebo městský písař; knihař a iluminátor až ve větším centru",
        "place": "Fara, hradní kancelář, městská správa, klášterní nebo městská dílna; přístup je společensky i jazykově omezený",
        "mods": "soustředění, paměť, pozorování, jemná obratnost, historie a dobový společenský protokol; gramotnost, latina, světlo a kvalita materiálu",
        "role": "Pozdní učené řemeslo. Hráč se neučí jen číst: připravuje inkoust a brk, opisuje, vede záznamy, skládá složky, šije knižní blok a opravuje vazbu. Otevírá Velkou knihu dílny.",
        "spec": "Opisování a kancelářské záznamy; knihařství a opravy; jednoduchá výzdoba, schémata a značky dílen.",
        "masterpiece": "Velká kniha dílny: přesný, svázaný a svědky potvrzený záznam postupů, vad a mistrovských výrobků všech třinácti oborů.",
        "rows": [
            ("Pozorovatel", "Psací materiály; výroba a seříznutí brku; inkoust; linkování; držení předlohy; význam pečetí a záznamů.", "Připraví pracoviště a přesně rozezná, co je opis, poznámka, účet, právní zápis a církevní text."),
            ("Učeň", "Tah písma; čísla a jednoduché účty; krátké opisy; sušení inkoustu; skládání a sešívání cvičných složek.", "Bez významové chyby opíše krátký zápis a vytvoří jednoduchý sešit, který se nerozpadá."),
            ("Tovaryš", "Delší texty; opravy chyb; rejstřík značek; výroba knižního bloku; desky, potah a jednoduché spony.", "Samostatně vytvoří použitelný pracovní zápisník a doloží jeho čitelnost, úplnost a odolnost."),
            ("Specialista", "Volba větve; úřední formule; latina podle původu postavy; náročné opravy vazby; schémata nástrojů a postupů.", "Zpracuje neúplnou nebo poškozenou předlohu, přizná nejistá místa a nezmění smysl textu."),
            ("Mistr", "Redakce rozsáhlého svazku; ověřování svědků; trvanlivé materiály; organizace opisovačů a řemeslníků vazby.", "Velká kniha dílny je úplná, dohledatelná, fyzicky odolná a přijata mistry ostatních oborů."),
            ("Inovátor / Mistr dílny", "Inovátor zlepšuje rejstřík, schémata nebo předávání pracovních znalostí. Mistr dílny vede opis, vazbu, materiál a kontrolu významu.", "Završení oboru schopností uchovat cizí zkušenost bez jejího zkreslení."),
        ],
    },
]

CRAFT_ORDER = [
    "smithing",
    "woodworking",
    "masonry",
    "wheelwrighting",
    "pottery",
    "leatherworking",
    "textiles",
    "bowyery",
    "brewing",
    "scribal",
    "herbalism",
    "healing",
    "cooking",
]
CRAFTS.sort(key=lambda craft: CRAFT_ORDER.index(craft["id"]))


def add_cover(doc):
    section = doc.sections[0]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HERNÍ NÁVRH")
    set_font(r, size=10.5, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Plánovaný strom řemesel")
    set_font(r, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Řemeslo středověku")
    set_font(r, size=16, color=DARK_BLUE, bold=True)

    strap = doc.add_paragraph()
    strap.paragraph_format.space_after = Pt(10)
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = strap.add_run("Pozoruj - pomoz - uč se - tvoř - předej dál")
    set_font(r, size=11, color=MUTED, italic=True)

    add_illustration(
        doc,
        ASSET_DIR / "artifact.png",
        title="Motyka a Velká kniha dílny",
        alt_text="Opravená středověká motyka, otevřená dílenská kniha a drobné symboly třinácti řemesel.",
        width=1.8,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("Návrh 2.0 | 11. srpna 2026")
    set_font(r, size=11, color=NAVY, bold=True)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2.add_run("Návaznost na Development Bible 4.0 a současný systém statistik hráče")
    set_font(r, size=9.5, color=MUTED, italic=True)

    doc.add_page_break()


def add_header_footer(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Řemeslo středověku | Strom řemesel")
        set_font(r, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("Strana ")
        set_font(r, size=9, color=MUTED)
        add_field(p, "PAGE")
        r = p.add_run(" z ")
        set_font(r, size=9, color=MUTED)
        add_field(p, "NUMPAGES")


def main():
    doc = Document()
    configure_styles(doc)
    bullet_num, decimal_nums = create_numbering(doc)
    add_header_footer(doc)
    add_cover(doc)

    doc.add_heading("1. Účel a pevná pravidla", level=1)
    add_callout(
        doc,
        "Hlavní rozhodnutí",
        "Strom není nákup schopností za zkušenostní body. Uzel se odemkne teprve tehdy, když hráč postup viděl, pochopil, prakticky provedl a prokázal na skutečné práci. Číselný postup slouží enginu a přehledu, ale sám o sobě nezaručí hodnost.",
        fill=ACCENT_FILL,
    )
    add_bullets(
        doc,
        [
            "Třináct hlavních světských řemesel má vlastní plný strom a vlastní důkazy zvládnutí; kratší pracovní dráhy dodávají suroviny, vztahy a místní zkušenost.",
            "Společné hodnosti jsou: Pozorovatel, Učeň, Tovaryš, Specialista, Mistr a závěrečná volba Inovátor / Mistr dílny.",
            "Hráč se učí od konkrétních lidí. Přístup k mistrovi, dílně a materiálu závisí na důvěře, pověsti a společenském postavení.",
            "Postupy se učí pozorováním a prací. Kniha může pomoci se znalostí, ale nenahradí motorickou dovednost ani mistrovu zkoušku.",
            "Chyba není jen trest. Rozpoznání vady, oprava a hospodárné zachránění výrobku jsou samostatné uzly stromu.",
            "Vyšší výrobky přirozeně propojují více oborů; hráč se tak nestane mistrem izolovaným opakováním jediného receptu.",
            "Skrytá větev návratu se odemyká až porozuměním hmotným řemeslům. Paranormální vrstva zůstává nejednoznačná, jak stanovuje vývojová bible.",
        ],
        bullet_num,
    )
    add_callout(
        doc,
        "Co znamená všechna řemesla",
        "Dokument zahrnuje všechna plánovaná hráčská řemesla projektu: 13 plných světských stromů, 6 kratších pracovních drah a skrytou větev artefaktu. Nejde o úplný soupis všech povolání známých v Českém království kolem roku 1400; tržiště a vzdálenější centra mohou dál nabízet specialisty, které hráč sám neovládá.",
    )

    doc.add_heading("2. Společný model postupu", level=1)
    rank_rows = [
        ("Pozorovatel", "Vidí práci, učí se názvy, bezpečnost, materiál a správné pořadí kroků.", "Bez samostatné výroby; smí asistovat a popisovat chyby."),
        ("Učeň", "Provádí základní operace pod dohledem a opakuje je v různých podmínkách.", "Jednoduché výrobky a opravy; mistr může zásah zastavit."),
        ("Tovaryš", "Zvládne celý běžný výrobek, plánuje postup a opravuje vlastní chyby.", "Samostatná práce a cesta za dalšími mistry; běžné zakázky."),
        ("Specialista", "Volí jednu hlubší větev a řeší neobvyklé vady, materiál a požadavky.", "Náročné zakázky; první mezioborové mistrovské projekty."),
        ("Mistr", "Navrhuje postup, ručí za jakost, hospodaří se zdroji a předává znalost.", "Mistrovské výrobky, posudky, vedení práce a přístup ke skrytým stopám."),
        ("Inovátor", "Zlepší nástroj, přípravek, konstrukci nebo pracovní postup v mezích doby.", "Technická závěrečná cesta; zlepšení musí prokázat užitek i spolehlivost."),
        ("Mistr dílny", "Řídí lidi, zásoby, termíny, bezpečnost a standard jakosti.", "Sociálně-organizační závěrečná cesta; umí vychovat dalšího pracovníka."),
    ]
    build_table(doc, ["Hodnost", "Co hráč skutečně umí", "Herní oprávnění"], rank_rows, [1500, 4460, 3400], font_size=9.1)

    doc.add_heading("2.1 Podmínky odemčení uzlu", level=2)
    add_numbers(
        doc,
        [
            "Znalost: hráč postup pozoroval, slyšel vysvětlení nebo jej správně odvodil.",
            "Praxe: provedl potřebné dílčí operace v požadovaném počtu a rozmanitosti.",
            "Jakost: alespoň několik výsledků dosáhlo stanovené úrovně bez skryté pomoci.",
            "Porucha a náprava: hráč poznal typickou vadu a alespoň jednou ji správně napravil nebo práci bezpečně zastavil.",
            "Přístup: má potřebný nástroj, pracoviště, materiál a svolení majitele či mistra.",
            "Svědectví: u hodnostních uzlů práci přijme mistr, objednavatel nebo komunita.",
        ],
        decimal_nums[0],
    )

    doc.add_heading("2.2 Stav znalosti postupu", level=2)
    recipe_rows = [
        ("Neznámý", "Hráč o postupu neví; předmět může vidět, ale neumí rozložit jeho výrobu."),
        ("Zaslechnutý", "Zná účel nebo název. Může se ptát, hledat mistra či suroviny."),
        ("Pozorovaný", "Viděl celý postup nebo jeho důležitou část. UI ukáže kroky, ne jistotu úspěchu."),
        ("Vyzkoušený", "Alespoň jeden pokus; stále se zobrazují široká okna a výrazná rizika."),
        ("Osvojený", "Postup opakovaně funguje v běžných podmínkách; hráč jej smí samostatně nabízet."),
        ("Zvládnutý", "Stabilní jakost, diagnostika vad a schopnost postup vysvětlit druhému."),
    ]
    build_table(doc, ["Stav", "Význam"], recipe_rows, [1800, 7560], font_size=9.3)

    doc.add_heading("2.3 Výpočet postupu oboru", level=2)
    add_labeled_paragraph(doc, "Doporučení", "Každý obor může mít interní hodnotu 0-100 %, ale hráči se mají zobrazovat hlavně konkrétní uzly, zkoušky a chybějící zkušenosti.")
    build_table(
        doc,
        ["Složka", "Váha", "Co se započítává"],
        [
            ("Znalost", "20 %", "Pozorované a pochopené postupy, materiály, bezpečnost a názvosloví."),
            ("Praktické operace", "35 %", "Rozmanité úkony, nikoli nekonečné opakování jednoho snadného výrobku."),
            ("Jakost výrobků", "20 %", "Stabilní výsledky, náročnost a práce bez skryté pomoci."),
            ("Vady a opravy", "15 %", "Diagnostika, záchrana materiálu, údržba nástrojů a bezpečné zastavení."),
            ("Svědectví a předání", "10 %", "Mistrova zkouška, přijatá zakázka, výuka pomocníka nebo prospěch komunitě."),
        ],
        [2100, 1100, 6160],
        font_size=9.2,
    )

    doc.add_heading("3. Přehled hlavních větví", level=1)
    overview_rows = [
        ("Kovářství", "Beneš", "Nástroje, podkovy, kování", "Nástrojař / podkovář / stavební kování", "Dřevo, kámen, kolářství"),
        ("Tesařství a truhlářství", "Matěj, tesař", "Stavby, nábytek, opravy", "Tesař / truhlář / stavební opravy", "Kovářství, kámen"),
        ("Kamenictví a zednictví", "Hradní mistr", "Zdi, ostění, pece", "Zedník / kameník / kamnář", "Dřevo, kov, hrnčířství"),
        ("Kolářství", "Kolář z Doksů", "Kola, nápravy, vozy", "Kola / vozy / vodní kola", "Dřevo, kov, mlýn"),
        ("Hrnčířství", "Hrnčíř", "Nádoby, zásoby, pece", "Nádobí / zásobnice / kachle", "Vaření, bylinky, kamenictví"),
        ("Koželužství a ševcovství", "Koželuh, švec", "Obuv, brašny, postroje", "Švec / řemenář / sedlář", "Textil, kov, dřevo"),
        ("Tkalcovství a krejčovství", "Grete, tkalec", "Plátno, oděv, opravy", "Tkalec / krejčí / barvíř", "Kůže, bylinky, ranhojičství"),
        ("Výroba luků a šípů", "Lesník, střelec", "Luky, šípy, tětivy", "Lukař / šípař / polní opravy", "Dřevo, kov, textil, kůže"),
        ("Pivovarnictví", "Hospodyně, pivovarník", "Slad, pivo, větší várky", "Sladování / var / provoz", "Mlýn, vaření, dřevo"),
        ("Písařství a knihařství", "Kněz, písař", "Záznamy, knihy, opravy", "Opis / vazba / schémata", "Všechna řemesla"),
        ("Bylinkářství", "Zdislava", "Určení, sběr, sušení, směsi", "Sběrač / léčivé směsi / lidová vrstva", "Ranhojičství, vaření"),
        ("Ranhojičství", "Zdislava, ranhojič", "Rány, kosti, nemocní", "Rány / kosti / nemoci", "Bylinky, textil, hrnčířství"),
        ("Vaření", "Anna", "Jídlo, chléb, zásoby", "Kuchyně / pečení / konzervace", "Hrnčířství, bylinky, mlýn"),
        ("Tajemství artefaktu", "Bez běžného mistra", "Návrat a smysl řemesel", "Hmota / paměť / společenství", "Všech 13 oborů"),
    ]
    build_table(doc, ["Větev", "První učitel", "Hlavní výstupy", "Specializace", "Vazby"], overview_rows, [1450, 1350, 2150, 2650, 1760], font_size=8.15)

    doc.add_heading("4. Detail stromů řemesel", level=1)
    intro = doc.add_paragraph("Každá níže uvedená větev používá společnou hodnostní osu, ale má vlastní znalosti, zkoušky a mistrovský projekt. Specializace se volí na úrovni Specialista; později lze otevřít další, avšak za cenu nové praxe a dalšího učitele. Ilustrace jsou orientační herní značky nástrojů a materiálů, nikoli tvrzení o jediném přesném vzhledu dobové dílny.")
    intro.paragraph_format.keep_together = True
    add_illustration(
        doc,
        ASSET_DIR / "main-sheet.png",
        title="Obrazový přehled hlavních řemesel a hospodářských drah",
        alt_text="Mřížka šestnácti ilustrací: kovářství, tesařství, kamenictví, kolářství, hrnčířství, koželužství, tkalcovství, výroba luků a šípů, pivovarnictví, písařství, bylinkářství, ranhojičství, vaření, zemědělství, mlynářství a rybářství.",
        width=4.65,
        caption="Obrazový rejstřík 13 hlavních řemesel a tří největších hospodářských drah.",
    )

    for idx, craft in enumerate(CRAFTS, start=1):
        doc.add_page_break()
        doc.add_heading(f"4.{idx} {craft['name']}", level=2)
        add_illustration(
            doc,
            ASSET_DIR / CRAFT_ICON_FILES[craft["id"]],
            title=f"Ilustrace oboru {craft['name']}",
            alt_text=CRAFT_ICON_ALTS[craft["id"]],
            width=1.12,
        )
        add_labeled_paragraph(doc, "Úloha ve světě", craft["role"])
        add_labeled_paragraph(doc, "První učitel", craft["mentor"])
        add_labeled_paragraph(doc, "Pracoviště", craft["place"])
        add_labeled_paragraph(doc, "Modifikátory", craft["mods"])
        build_table(doc, ["Hodnost", "Uzly a schopnosti", "Důkaz zvládnutí"], craft["rows"], [1200, 5050, 3110], font_size=8.85)
        add_labeled_paragraph(doc, "Specializace", craft["spec"])
        add_labeled_paragraph(doc, "Mistrovský projekt", craft["masterpiece"])

    doc.add_page_break()
    doc.add_heading(f"4.{len(CRAFTS) + 1} Tajemství artefaktu - skrytá větev návratu", level=2)
    add_illustration(
        doc,
        ASSET_DIR / "artifact.png",
        title="Motyka, Velká kniha dílny a svědectví řemesel",
        alt_text="Opravená motyka vedle otevřené knihy a třinácti drobných symbolů řemesel.",
        width=1.55,
    )
    add_callout(
        doc,
        "V UI se nemusí nazývat magie",
        "Tato větev může být pro hráče nejprve vedena jako Tajemství artefaktu, Paměť místa nebo Poznání motyky. Zachová se tím pravidlo vývojové bible: jevy nikdy nejsou úplně potvrzené a mohou mít duchovní, psychologické i nadpřirozené vysvětlení.",
        fill=CAUTION_FILL,
    )
    hidden_rows = [
        ("První ozvěna", "3 světské obory alespoň 40 %", "Artefakt reaguje na poctivě dokončený výrobek, místo nebo sen. Hráč ještě neví proč."),
        ("Spojení stop", "6 oborů alespoň 60 %", "V materiálech, značkách mistrů, lidových zvycích a vyprávění se objeví společný vzorec."),
        ("Paměť práce", "4 obory na hodnosti Mistr", "Mistrovské výrobky odemykají části minulosti nástroje a vazbu na Matějovu rodinu."),
        ("Třináct svědectví", "Všech 13 světských oborů alespoň 80 %", "Každý obor dodá ověřený výrobek, postup, opravu nebo svědectví pro závěrečnou práci."),
        ("Velká kniha dílny", "Písařství a knihařství: Mistr", "Hráč sváže a zpřístupní ucelený záznam postupů, vad, mistrů a důkazů všech třinácti oborů."),
        ("Ztracený rukopis", "Výzkum Housky a důvěra duchovních či správy", "Nalezený neúplný text propojí artefakt, místo a rodinnou paměť, ale nedá jednoznačný návod ke kouzlu."),
        ("Obnova artefaktu", "Mezioborový mistrovský projekt", "Hráč obnoví motyku a připraví místo, čas, nádobu, oděv, léčebné prostředky a svědky."),
        ("Návrat / setrvání", "Dokončený příběhový uzel", "Hráč provede závěrečný čin na Housce. Výsledek může být návrat, odmítnutí návratu nebo nejednoznačný konec."),
    ]
    build_table(doc, ["Uzel", "Podmínka", "Obsah"], hidden_rows, [1850, 2600, 4910], font_size=9.0)
    add_labeled_paragraph(doc, "Tři závěrečné směry", "Hmota - přesnost výrobků a obnova artefaktu; Paměť - historie, symboly, sny a místa; Společenství - předání řemesel, důvěra a pomoc lidí. Směr mění podobu finále, ne požadavek na zvládnutí světských oborů.")

    doc.add_heading("5. Mezioborové uzly", level=1)
    cross_rows = [
        ("Vozové kolo s obručí", "Kolářství: Tovaryš; kov: Učeň", "Náboj, paprsky, loukotě, obruč a souosost", "Kolář / kovář"),
        ("Truhla s kováním", "Truhlářství: Tovaryš; Kov: Učeň", "Přesný korpus, panty, petlice a případně zámek", "Truhlář / kovář"),
        ("Pec a kamna", "Kamenictví: Tovaryš; Hrnčířství: Tovaryš", "Tah, tepelná roztažnost, kachle a bezpečný provoz", "Zedník / hrnčíř"),
        ("Léčebná souprava", "Bylinky, ranhojičství a textil: Tovaryš", "Obvazy, nádoby, sušené směsi, popisy a bezpečné dávky", "Zdislava / ranhojič"),
        ("Postroj a vůz", "Kůže: Specialista; Dřevo: Tovaryš; Kov: Tovaryš", "Přenos síly, pohodlí zvířete, řemení a opravy", "Sedlář / kolář"),
        ("Lovecká souprava", "Luky a šípy: Tovaryš; kov, textil a kůže: Učeň", "Hroty, tětiva, toulec, bezpečné sladění luku a šípů", "Střelec / lukař"),
        ("Várka pro slavnost", "Pivo: Specialista; mlýn a vaření: rozvinutá dráha", "Slad, voda, palivo, nádoby, čistota, čas a výdej", "Pivovarník / rychtář"),
        ("Hradní dveře", "Dřevo, kov a kamenictví: Specialista", "Rám, křídlo, závěsy, petlice, ostění a osazení", "Hradní mistr"),
        ("Zásoby na zimu", "Vaření: Specialista; hrnčířství a bylinky: Učeň", "Sušení, solení, kvašení, nádoby a evidence zásob", "Anna / mlynář"),
        ("Velká kniha dílny", "Písařství: Mistr; všech 13 oborů má svědectví", "Ověřené postupy, vady, schémata, rejstřík a trvanlivá vazba", "Mistři oborů / písař"),
        ("Obnova motyky", "Kov a dřevo: Mistr; všech 13 oborů alespoň 80 %", "Fyzické jádro závěrečné obnovy artefaktu", "Bez jednoho mistra; kolektivní projekt"),
    ]
    build_table(doc, ["Projekt", "Předpoklady", "Co prověřuje", "Kdo posuzuje"], cross_rows, [1900, 2600, 3160, 1700], font_size=8.8)

    doc.add_page_break()
    doc.add_heading("6. Vedlejší pracovní větve", level=1)
    add_labeled_paragraph(doc, "Důvod", "Ne každá historická práce musí být plnohodnotný strom nutný k návratu. Některé činnosti lépe fungují jako kratší profesní dráhy, které dodávají suroviny, vztahy a místní znalost.")
    side_rows = [
        {
            "name": "Zemědělství",
            "image": "main-13.png",
            "alt": "Dřevěný pluh, srp a snop obilí.",
            "activities": "Orba motykou a pluhem, setí, seno, žně, mlácení a skladování.",
            "links": "Vaření, tkalcovství, pivovarnictví, kovářství a chov zvířat.",
            "scope": "Plnější sezónní dráha po rozšíření polí; nepatří mezi 13 povinných mistrovství.",
        },
        {
            "name": "Mlynářství",
            "image": "main-14.png",
            "alt": "Mlýnský kámen, lopatka na zrní a vodní kolo.",
            "activities": "Pytle, tok vody, čištění kamene, mletí, mouka a základní údržba převodů.",
            "links": "Vaření, pivovarnictví, kolářství, dřevo, kov a kamenictví.",
            "scope": "Silná profesní dráha navázaná na mlýn a Ondřeje; vyšší opravy jsou mezioborové.",
        },
        {
            "name": "Rybářství a rybníkářství",
            "image": "main-15.png",
            "alt": "Proutěná vrš, rybářská síť, ryba a dřevěné stavidlo.",
            "activities": "Sítě, vrše, čluny, třídění úlovku, břehy, hráze, stavidla a hospodaření s rybníkem.",
            "links": "Provaznictví, dřevo, vaření a obchod v Doksech.",
            "scope": "Otevřít po zapojení Velkého rybníka a Doks; odbornější rybníkářství vyžaduje místního správce.",
        },
        {
            "name": "Chov zvířat",
            "image": "side-00.png",
            "alt": "Dřevěné jho, zvonec, nůžky na ovčí vlnu a píce.",
            "activities": "Pasení, krmení, dojení, stříž, čištění stáje, ošetření a rozpoznání potíží.",
            "links": "Kůže, textil, vaření, zemědělství a kolářství.",
            "scope": "Postup přes důvěru domácnosti a zodpovědnost za živé zvíře, ne přes výrobu předmětů.",
        },
        {
            "name": "Lesní práce a uhlířství",
            "image": "side-01.png",
            "alt": "Lesní sekera, polena, dřevěné uhlí a zemní milíř.",
            "activities": "Výběr dřeva, kácení, štípání, klestí, doprava a pomalé pálení milíře.",
            "links": "Kovářství, tesařství, hrnčířství, pivovarnictví a vaření.",
            "scope": "Riziková surovinová dráha s lesními právy a právními následky; železo se do oblasti dováží.",
        },
        {
            "name": "Provaznictví a košíkářství",
            "image": "side-02.png",
            "alt": "Svinutý konopný provaz, stáčecí hák, vlákna a proutěný koš.",
            "activities": "Příprava konopí a lýka, stáčení provazů, uzly, sítě, vrše, koše a jednoduché opravy.",
            "links": "Textil, rybářství, zemědělství, dřevo, doprava a domácnost.",
            "scope": "Krátká praktická dráha; vyšší lana a sítě vyžadují rovnoměrnost, měření a zkoušku zatížením.",
        },
    ]
    build_illustrated_side_table(doc, side_rows)

    doc.add_heading("7. Jakost, vady a podmínky práce", level=1)
    quality_rows = [
        ("Zkažený", "Výrobek neplní funkci nebo je nebezpečný.", "Ztráta materiálu, možné zranění, mistr zasáhne.", "Prasklá nádoba; přepálené železo; plesnivá bylina."),
        ("Provizorní", "Funguje krátce nebo jen v mírných podmínkách.", "Nízká odolnost, horší cena, nouzové použití.", "Hrubá záplata; měkké topůrko; řídká kaše."),
        ("Běžný", "Odpovídá každodenní potřebě a přiměřeným nákladům.", "Standardní výkon, cena a životnost.", "Hrnec, opasek nebo nástroj pro běžnou domácnost."),
        ("Poctivý", "Nadprůměrně přesný, trvanlivý nebo úsporný.", "Lepší důvěra, delší životnost, doporučení mistra.", "Rovnoměrné plátno; dobře vyvážená sekera."),
        ("Výborný", "Náročná práce s minimem vad a dobrým využitím materiálu.", "Přístup k prestižním zakázkám a mistrovským uzlům.", "Hradní dveře; kvalitní postroj; zásoba léčiv."),
        ("Mistrovský", "Opakovatelná špičková jakost a správné řešení skutečné potřeby.", "Doklad hodnosti Mistr; příběhové a společenské důsledky.", "Projekt přijatý dalšími mistry nebo komunitou."),
    ]
    build_table(doc, ["Stupeň", "Význam", "Následek", "Příklad"], quality_rows, [1300, 2600, 2710, 2750], font_size=8.65)
    doc.add_heading("7.1 Doporučený model výsledku", level=2)
    add_bullets(
        doc,
        [
            "Materiál: vhodnost, čistota, vlhkost, stáří a skryté vady.",
            "Postup: správné pořadí, načasování, teplota, síla a přesnost jednotlivých kroků.",
            "Nástroj a pracoviště: opotřebení, ostrost, stabilita, světlo, počasí a dostupná pomoc.",
            "Stav hráče: únava, hlad, bolest, stres, zima a mokro ovlivňují ovládání i rozhodování.",
            "Dovednost: hodnost rozšiřuje toleranci a diagnostiku, ale neodstraňuje nutnost správně hrát minihru.",
            "Náprava: včas poznaná chyba může snížit jakost, ale zachránit výrobek a přidat zkušenost s vadou.",
        ],
        bullet_num,
    )

    doc.add_heading("8. Vliv zvoleného původu postavy", level=1)
    profile_rows = [
        ("Zálesák", "Dřevo, oheň, lana, improvizace, sběr a pohodlí v přírodě.", "Rychleji pochopí praktický smysl kroků; nedostává automaticky středověké recepty ani mistrovskou přesnost."),
        ("Student historie", "Pozorování, paměť, soustředění, historie, symboly a část dobového názvosloví.", "Snáz odvodí účel nástroje a zapamatuje postup; motorickou práci musí stále nacvičit."),
        ("Měšťan", "Empatie, etiketa, vyjednávání, společenský protokol a práce mezi lidmi.", "Rychleji získá učitele, zakázku nebo materiál; slabší praktický začátek vyvažuje sociálním přístupem."),
    ]
    build_table(doc, ["Původ", "Výchozí výhody", "Zásadní omezení"], profile_rows, [1800, 3000, 4560], font_size=9.0)
    add_callout(doc, "Pravidlo rovnosti cest", "Žádný původ nesmí uzamknout obor ani přeskočit hodnostní zkoušku. Původ mění způsob učení a první příležitosti, nikoli konečný strop.")

    doc.add_heading("9. Návrh uložení v enginu", level=1)
    add_labeled_paragraph(doc, "Architektonické rozhodnutí", "Stávající PlayerStats.h už dobře obsahuje tělesné, mentální, sociální a obecné praktické modifikátory. Trvalý stav stromu je vhodné držet v samostatném CraftProgressionState, aby se nemíchala kondice postavy s odemčenými uzly, postupy a mistrovskými zkouškami.")
    data_rows = [
        ("CraftDomain", "smithing, woodworking, masonry, wheelwrighting, pottery, leatherworking, textiles, bowyery, brewing, scribal, herbalism, healing, cooking, artifact"),
        ("WorkTrackDomain", "farming, milling, fishing, husbandry, forestryCharcoal, ropeBasketry"),
        ("CraftRank", "Observer, Apprentice, Journeyman, Specialist, Master, Innovator, WorkshopMaster"),
        ("CraftNodeDef", "id, domain, rank, prerequisites, requiredActions, qualityGate, mentorGate, toolGate, unlocks"),
        ("CraftNodeState", "unknown, revealed, observed, practiced, completed; counters a nejlepší dosažená jakost"),
        ("ProcedureKnowledge", "heard, observed, attempted, stable, mastered"),
        ("CraftEvidence", "výrobek, oprava, odhalená vada, svědek, datum, pracoviště a použité podmínky"),
        ("SpecializationState", "zvolená hlavní specializace, vedlejší otevřené větve, mistr a splněná zkouška"),
    ]
    build_table(doc, ["Datový celek", "Obsah"], data_rows, [2500, 6860], font_size=9.0)

    doc.add_heading("9.1 Vyhodnocení uzlu", level=2)
    add_numbers(
        doc,
        [
            "Ověřit předpoklady a zda je uzel hráči vůbec známý.",
            "Ověřit dostupnost mistra, pracoviště, nástroje a materiálu.",
            "Sečíst pouze platné důkazy; stejné snadné opakování má klesající přínos.",
            "Ověřit minimální jakost, samostatnost a případnou diagnostiku vady.",
            "U hodnostního uzlu vyžádat svědectví NPC nebo dokončení komunitního projektu.",
            "Po odemčení zapsat nové činnosti, postupy, výrobky, opravy a dialogové možnosti.",
        ],
        decimal_nums[1],
    )

    doc.add_heading("10. Doporučené pořadí realizace", level=1)
    phase_rows = [
        ("1. Svislý řez", "Bylinkářství do Tovaryše", "Navazuje na hotové určování rostlin, sběr, inventář, statistiky a Zdislavu. Ověří obecný datový model uzlů."),
        ("2. Domácnost", "Vaření do Učně", "Propojí oheň, nádoby, suroviny, čas, stav hráče a přijetí u Matějovy rodiny."),
        ("3. Opravy", "Dřevo do Učně", "Plot, topůrko a vrata vytvoří rychlé, viditelné výsledky a první práci s vadami."),
        ("4. Dílna", "Kovářství do Učně", "Pomoc u měchů, hřebíky a oprava srpu ověří mentora, pracoviště a teplotní minihru."),
        ("5. Společný projekt", "Léčebná souprava", "První mezioborový uzel propojí bylinky, ranhojičství, textil a nádoby."),
        ("6. Rozšíření světa", "Textil, kůže, hrnčířství, kámen", "Přidá městské a hradní učitele, obchodní vazby a dlouhé výrobní řetězce."),
        ("7. Doprava a lov", "Kolářství; luky a šípy", "Propojí dílny s cestami, Bezdězem, lesem, kovářstvím a ověřováním výrobku v pohybu."),
        ("8. Hospodářství", "Pivovarnictví a kratší pracovní dráhy", "Zapojí mlýn, pole, rybníky, chov, lesní práva, sezóny a místní trh."),
        ("9. Paměť řemesel", "Písařství a Velká kniha dílny", "Pozdní učená větev spojí důkazy všech oborů a připraví podmínky ztraceného rukopisu."),
        ("10. Pozdní hra", "Mistrovství a artefakt", "Teprve po stabilní ekonomice, důvěře NPC a víceoborových projektech otevřít návratovou větev."),
    ]
    build_table(doc, ["Etapa", "Rozsah", "Proč právě tehdy"], phase_rows, [1700, 2350, 5310], font_size=9.0)

    doc.add_heading("11. Kritéria hotového prvního stromu", level=1)
    add_bullets(
        doc,
        [
            "Každý uzel má jasný herní úkon, podmínku, výsledek a způsob ověření.",
            "Alespoň jedna chyba je rozpoznatelná během práce a jedna až po dokončení.",
            "Mentor reaguje jinak na neznalost, nepozornost, poctivý neúspěch a úspěšnou nápravu.",
            "Postup se uloží a načte bez ztráty důkazů, specializace nebo stavu známých postupů.",
            "Statistiky hráče modifikují obtížnost, ale samy neodemknou uzel.",
            "Mistrovský výrobek má použití ve světě a viditelný dopad na lidi, ne pouze vyšší prodejní cenu.",
            "První plná větev je znovupoužitelný vzor pro ostatní obory, nikoli jednorázově naprogramovaná výjimka.",
        ],
        bullet_num,
    )

    doc.add_heading("12. Podklady a vazba na současný projekt", level=1)
    add_bullets(
        doc,
        [
            "Development Bible 4.0 - Řemeslo války: region Houska kolem roku 1400, společenská integrace, herbář, minihry, každodenní práce a nejednoznačná paranormální vrstva.",
            "Regionální pravidlo: Houska - Doksy - Velký rybník - Bezděz - Podbezdězí stojí na specializaci, trhu a dovozu. Železo a hotové polotovary se dovážejí; dokument nepředpokládá neprokázaný místní železný důl kolem roku 1400.",
            "PlayerStats.h: obecné atributy, survival dovednosti, craft.toolRepair, clothingRepair, ropeWork, woodcraft, improvisation a mentální hodnoty focus, memory a observation.",
            "Dosavadní projektové rozhodnutí: návrat je možný až po zvládnutí všech 13 hlavních světských řemesel alespoň na 80 %, vytvoření Velké knihy dílny, nalezení ztraceného rukopisu a dokončení skryté větve artefaktu.",
            "Dosavadní projektové rozhodnutí: postup je Pozorovatel -> Učeň -> Tovaryš -> Specialista -> Mistr -> Inovátor / Mistr dílny.",
        ],
        bullet_num,
    )

    props = doc.core_properties
    props.title = "Plánovaný strom řemesel - Řemeslo středověku"
    props.subject = "Herní návrh postupu řemesel, specializací, jakosti a návratové větve"
    props.author = "Radek Lhoták"
    props.keywords = "Řemeslo středověku, skill tree, řemesla, Houska, game design"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
